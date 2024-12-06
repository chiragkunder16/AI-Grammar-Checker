#Final code for report generation.

import os
import streamlit as st
from groq import Groq
from fpdf import FPDF
from docx import Document

# Initialize the Groq client with the API key
client = Groq(
    api_key=os.environ.get("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"),
)

def check_grammar(input_text):
    """
    Function to check grammar using Groq's LLaMA model.
    Takes an input text with potential grammatical errors and returns the corrected text.
    """
    prompt = f"Correct the grammar in the following sentence: '{input_text}'."
   
    # Make the API request to LLaMA model
    chat_completion = client.chat.completions.create(
        messages=[{
            "role": "user",
            "content": prompt,
        }],
        model="llama3-8b-8192",  # Model to use
    )
 
    # Get and return the corrected sentence from the response
    return chat_completion.choices[0].message.content

def change_tone(input_text, tone):
    """
    Function to change the tone of the sentence (Formal or Friendly) after grammar correction.
    """
    if tone == "Formal":
        prompt = f"Make the following sentence sound more formal: '{input_text}'."
    else:
        prompt = f"Make the following sentence sound more friendly: '{input_text}'."
   
    # Make the API request to LLaMA model
    chat_completion = client.chat.completions.create(
        messages=[{
            "role": "user",
            "content": prompt,
        }],
        model="llama3-8b-8192",  # Model to use
    )
    
    # Return the tone-adjusted sentence
    return chat_completion.choices[0].message.content

def generate_pdf_and_word(original, corrected, tone_changed, tone):
    """
    Generate both a PDF and a Word report containing the original sentence, corrected sentence, and the tone change.
    """
    # Generate PDF
    pdf = FPDF()
    pdf.add_page()
    
    # Set font for title and headings
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt="Grammar Correction and Tone Adjustment Report", ln=True, align="C")
    pdf.ln(10)  # Add line break
    
    # Set font for section headers
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="1. Original Sentence", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, f"{original}")
    pdf.ln(5)
    
    # Set font for section headers
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="2. Corrected Sentence", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, f"{corrected}")
    pdf.ln(5)
    
    # Set font for section headers
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="3. Tone Change", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, f"Tone changed to: {tone}")
    pdf.ln(5)
    
    # Add mistakes corrected section (you can expand this in the future)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="4. Mistakes Corrected", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, "Mistakes Corrected: (This will be more detailed in the future)")
    pdf.ln(5)
    
    # Define the directory path for saving the PDF and Word files
    report_directory = r"C:\Users\jkunder\Desktop\Gen ai\Reports"
    
    # Ensure the directory exists
    if not os.path.exists(report_directory):
        os.makedirs(report_directory)
    
    # Save PDF file to the specified directory
    pdf_output_path = os.path.join(report_directory, "grammar_tone_report.pdf")
    pdf.output(pdf_output_path)

    # Generate Word document
    doc = Document()
    doc.add_heading("Grammar Correction and Tone Adjustment Report", 0)
    doc.add_paragraph(f"Original Sentence: {original}")
    doc.add_paragraph(f"Corrected Sentence: {corrected}")
    doc.add_paragraph(f"Tone Changed to: {tone}")
    doc.add_paragraph("Mistakes Corrected: (This will be more detailed in the future)")
    
    # Save Word file to the same directory
    word_output_path = os.path.join(report_directory, "grammar_tone_report.docx")
    doc.save(word_output_path)

    return pdf_output_path, word_output_path

# Streamlit UI setup
st.title("AI Grammar Checker with Tone Adjustment")
st.write("Enter a sentence with possible grammar errors, correct it, adjust the tone if needed, and generate a report.")

# Text input field for the user to enter a sentence
input_sentence = st.text_area("Enter sentence to check:", height=150)

# Button to check grammar
if st.button("Correct Grammar"):
    if input_sentence:
        corrected_sentence = check_grammar(input_sentence)
        st.subheader("Corrected Sentence:")
        st.write(corrected_sentence)
        st.session_state.corrected_sentence = corrected_sentence  # Store for later use
        st.session_state.tone_change_requested = False  # Reset tone change flag
    else:
        st.warning("Please enter a sentence to check.")

# Ask if user wants to change the tone, after grammar correction
if 'corrected_sentence' in st.session_state:
    tone_change_option = st.radio("Would you like to change the tone of the corrected sentence?", ["Yes", "No"], index=1)  # Default is "No"
    
    if tone_change_option == "Yes":
        tone_option = st.selectbox("Select the tone for the corrected sentence:", ["Formal", "Friendly"])
        
        # Button to generate the tone-changed sentence
        if st.button("Generate Tone Change"):
            corrected_sentence = st.session_state.corrected_sentence
            if corrected_sentence:
                tone_changed_sentence = change_tone(corrected_sentence, tone_option)
                st.subheader(f"Tone Adjusted Sentence ({tone_option}):")
                st.write(tone_changed_sentence)
                st.session_state.tone_changed_sentence = tone_changed_sentence  # Store for report
                st.session_state.tone_change_requested = True  # Indicate that tone change was applied
            else:
                st.warning("Please correct the grammar first.")
    elif tone_change_option == "No":
        # No tone change, directly generate report
        st.session_state.tone_changed_sentence = st.session_state.corrected_sentence
        st.session_state.tone_change_requested = False

# Button to generate PDF and Word report (only available if the tone was changed or not)
if 'tone_changed_sentence' in st.session_state:
    if st.button("Generate Report"):
        original = input_sentence
        corrected = st.session_state.corrected_sentence
        tone_changed = st.session_state.tone_changed_sentence
        tone = "No change" if not st.session_state.tone_change_requested else st.session_state.tone_changed_sentence

        pdf_path, word_path = generate_pdf_and_word(original, corrected, tone_changed, tone)
        
        st.success("Report generated successfully!")
        
        # Provide buttons for downloading the generated files
        st.download_button(
            label="Download PDF Report",
            data=open(pdf_path, "rb").read(),
            file_name="grammar_tone_report.pdf",
            mime="application/pdf"
        )
        
        st.download_button(
            label="Download Word Report",
            data=open(word_path, "rb").read(),
            file_name="grammar_tone_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )